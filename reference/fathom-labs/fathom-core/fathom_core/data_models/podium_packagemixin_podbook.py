import fathom_core as core
from .system_flag import SystemFlag

from orator import Model
from orator.orm import has_one, belongs_to, has_many, has_many_through, scope
from orator import accessor, mutator
from orator.query.join_clause import JoinClause

import boto3
from botocore.errorfactory import ClientError
from rev_ai import apiclient
from deepgram import Deepgram
from bs4 import BeautifulSoup
import time
import asyncio
import requests

from pathlib import Path
import shutil
import uuid
import json
import datetime
import traceback
import os
from zipfile import ZipFile
from contextlib import contextmanager
import functools
import re
import pprint

from sendgrid.helpers.mail import *
from sendgrid import SendGridAPIClient

import numpy as np

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, ns

from num2words import num2words

class PodiumPackagePodbookMixin:
    
    def can_generate_podbook(self):
      return self.has_stored_transcription() \
          and self.has_generated_chapters() \
          and self.duration is not None \
          and self.duration > PodiumPackage.MIN_VECTORIZE_DURATION

    def has_generated_podbook(self):
        return self.podium_package_processing_configuration.generate_podbook_status == 'complete'

    def should_generate_podbook(self):
        return self.podium_package_processing_configuration.generate_podbook \
            and not self.has_generated_podbook()

    def generate_podbook(self, force=False):
        if self.should_generate_podbook() or force:
            try:
                self.remove_podbook_assets()
                self.podium_package_processing_configuration.generate_podbook_status = 'processing'
                self.podium_package_processing_configuration.save()
                self.generate_podbook_assets()
                self.podium_package_processing_configuration.generate_podbook_status = 'complete'
                self.podium_package_processing_configuration.save()
            except Exception as e:
              self.podium_package_processing_configuration.generate_podbook_status = 'error'
              self.podium_package_processing_configuration.save()
              raise e
            
    def remove_podbook_assets(self):
        core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'podbook_episode_chapter') \
            .delete()
        
        core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'podbook_episode_chapter_notes') \
            .delete()
            
    def generate_podbook_assets(self):
        package_title = self.get_accepted_asset('titles').content
        if self.podcast_episode_id:
            podcast_episode = core.data_models.PodcastEpisode.find(self.podcast_episode_id)
            if podcast_episode and podcast_episode.title:
                package_title = podcast_episode.title

        package_summary = self.get_accepted_asset('show_notes_summary').content

        system_message = f"""
You are Malcolm Gladwell ghostwriting a book for a podcast host.
You always write in the voice of the podcast host(s).
You never mention or reference the "section" or "sub-section" in your writing.
You never use the phrase "Reference Podcast Episode Source Content" in your writing.
You never start paragraphs with 'In conclusion' or 'In summary', or in similar cliche ways.
You never use the words 'dive', 'diving', or 'delve' as flourishes in your writing.
You never mention the podcast, you never mention the podcast host, and you never mention the fact that the source content for the book are podcast transcripts.
Importantly, you write in a way that mimics the podcast host's tone.
When a previous section is present, you never re-introduce the conversation with language like 'in a recent discussion' or 'in a recent conversation' - instead you just use the guests name when referring to things they said.
"""
        
        transcript = self.podium_package_transcript_file()
        
        transcript.load_content()
        sentences = transcript.get_sentences()
        
        last_section_title = ''
        last_section_content = ''

        for index, db_chapter in enumerate(self.chapters):
            chapter_sentences = core.data_models.PodiumPackageTranscriptFile.get_sentences_between_start_end(sentences, db_chapter.start, db_chapter.end)
            chapter_text = self.get_podbook_transcript_text_with_speakers_from_sentences(chapter_sentences)
            
            previously_written_instructions = "\nThe previously written section from the chapter has been included for your reference and to aid with crafting an appropriate first paragraph for this new section with a smooth transition. Notice that the conversation has already been introduced, and that you can simply refer to any guests by name if there are any guests. Remember that you are writing for the transcript above. Refrain from being overly repetitive.\n"

            transition_thoughts = "\n[Transition Thoughts: I'll be sure to transition from the previous section in the my opening paragraph to keep the flow of the writing smooth...Importantly, based on the previous section the conversation has already been introduced so I'll assume the reader is aware that I had a conversation with the guest - I can simply refer to the guest by name without re-introducing their title or the conversation.]\n"
            
            thought_instructions = "- Importantly, start your writing with thoughts about the names and roles of each speaker in each episode, followed by thoughts about the style and tone of the podcaster's voice in the podcast episode transcript, then adding thoughts about transitioning from the previous section, then by your thoughts about any formatted lists to include, and finally adding thoughts about the perspective you are writing from so you do not make a mistake."
            
            previous_section = f"""
[START previous chapter section you wrote]
SECTION TITLE: {last_section_title}
{last_section_content}
[END previous chapter section you wrote]
"""

            is_first_chapter = index == 0
            if is_first_chapter:
                previous_section = ""
                previously_written_instructions = ""
                transition_thoughts = ""
                thought_instructions = "- Importantly, start your writing with thoughts about the names and roles of each speaker in each episode, followed by thoughts about the style and tone of the podcaster's voice in the podcast episode transcript, then by your thoughts about any formatted lists to include, and finally adding thoughts about the perspective you are writing from so you do not make a mistake."

            episode_chapter_prompt = f"""
{previous_section}

[START podcast episode information]

Episode title: 
{package_title}

Summary of entire the episode:
{package_summary}

[END podcast episode information]

[START podcast episode transcript portion for the new section you are writing]
{chapter_text}
"""
                
            instruction_prompt = f"""
[END podcast episode transcript portion for the new section you are writing]

In this task, you are acting as a ghostwriter converting a section of a podcast episode transcript into one chapter sub-section for a written book. The book should be written in a manner that closely matches the style, tone, and voice of the podcast host. Consider the flow and structure of a book chapter section while maintaining the integrity of the original content. 
{previously_written_instructions}

The transcribed dialogue from the podcast will serve as the basis of this written work. Narratives should be adapted and transformed to read like written literature rather than spoken dialogue. Please apply proper grammar, punctuation, and sentence structure while ensuring the host's unique characteristics are maintained. Make sure to include any major points, stories, anecdotes, or arguments made during the podcast into the text in an engaging, cohesive manner. Remember to include any details that are likely to educate and inform the reader, giving them new ideas and approaches to problems they may have.

Depending on the context, consider restructuring Q&A portions into a more narrative or explanatory format, making it read more like a story or a lecture rather than a conversation. Feel free to use lists to break up the text for easy reading and highlight important points or processes.

The ultimate aim of this task is to create a readable, engaging, and coherent book chapter section reflective of the podcast host's voice and the content of the podcast episode. The audience should be able to read this text and hear the podcast host's voice in their head as if they are listening to the podcast episode.

The podcast episode transcript may indicate changes in speakers, but does not mention the actual name of the speaker or their role - like host or guest. So, you'll need to do your best to figure out which of the speakers is the host and which are guests in order to write from the correct perspective. Often times, the host will be Speaker 1, but not always. Guest names are typically mentioned in the episode information. Sometimes there is no guest, just the podcast host.

Follow these rules:

- Write one section. Per the example format below, write a single section title.

- Write in the first person from the perspective of the podcast host. Use "I" and "me" when writing. When framing the stories and lessons of guests, use their name(s) and "they" or "them".

- Importantly, focus on relating stories and explanations. Provide specific details and examples to clearly illustrate the points being made. Avoid writing in a way that is too general or vague. 

- Critically, it is imperative to write from the correct perspective. If there are multiple speakers, there may be errors in the speaker labels and a bit of overlap (the guest speaker's words appearing as if they are the hosts and vice versa) - you must correct for this. Always consider that you are writing from the host's perspective when relating stories told by the guest or guests.

- Refrain from writing narrative based on the podcast episode information. Its purpose is to provide you accurate information about any guests. Only use the transcript portion of the episode you were provided.

- Address the reader directly: Use second-person language like "you" and "your" to involve readers in the content and make the reading experience more engaging.

- Follow along with the podcast episode transcript, using facts and analogies as presented. Ignore all ads and sponsorships. You may need occasionally to fill in factual information to create a nice flow - but, don't invent or embellish any stories. Also, sometimes the transcript will have errors; use your expert knowledge to correct them.

- The sub-section you are working on is not the last sub-section of the chapter, so, refrain from writing a concluding last paragraph.

- Write as many paragraphs as needed to cover the transcript.

- Open each paragraph in a different way - refrain from over-repetition.

- Pay close attention to your own Style and Tone Thoughts.

- Use lists to break the content up and make it easy to read and understand.

- If the transcript mostly contains ads or sponsorships, simply write 'IGNORE'

{thought_instructions}

Please use this example format (it's very important to not forget the opening and closing brackets): 

[Speaker Thoughts: In episode "The Nature of Reality" Speaker 1 seems to be the podcast host who I am writing for. Speaker 2 seems to be the guest, John Foster. It's critical to write from the perspective of the podcast host while focusing on relating the stories of the guest(s), so I'll pay close attention to the transcript and correct for any speaker errors]

[Style and Tone Thoughts: The podcaster said these interesting phrases, '...', '...', and...with a tone and style like...we'll be sure to use these phrases with a style that sounds like the podcast host as if I were them...]
{transition_thoughts}
[List Thoughts: There are several potential lists of items, but they are not tips, tricks, key points or major take-aways so we'll keep them as narrative...the list of tips the guest offers, however, would benefit the reader if presented in a formatted list...]

[Perspective Thoughts: I am writing from the host's perspective and will relate any guest stories as their history, not mine...I will remember to include any details that are likely to educate and inform the reader, giving them new ideas...]

SECTION TITLE: ... 
...
"""
            max_response_length = 1600
            prompt_token_count = core.text.count_number_of_tokens(instruction_prompt, model='gpt-4o')
            remaining_tokens = 8150 - max_response_length - prompt_token_count
            trimmed_episode_chapter_prompt = episode_chapter_prompt.split(' ')[:int(remaining_tokens / 1.3)]
            episode_chapter_prompt = " ".join(trimmed_episode_chapter_prompt)

            prompt = episode_chapter_prompt + instruction_prompt

            print(prompt)
            print("******************************************************")
            raw_chapter_content = core.inference.gpt_chat_api_single_prompt(prompt, 'gpt-4o', system_message=system_message, max_tokens=max_response_length, temperature=1.2, top_p=0.8)
            #time.sleep(15)
            print("RAW CHAPTER CONTENT:")
            print(raw_chapter_content)

            if raw_chapter_content == "IGNORE":
                print("******************************************************")
                print("IGNORED")
                print("******************************************************")
                continue

            chapter_content = self.remove_podbook_thoughts(raw_chapter_content)
            title, content = self.get_podbook_title_and_content(chapter_content)
            print("******************************************************")
            print("FINAL CHAPTER CONTENT:")
            print(f"{title}")
            print(f"{content}")
            print("******************************************************")

            last_section_title = title
            last_section_content = content

            content_notes = self.generate_podbook_episode_chapter_notes(content)
        
            derived_from_asset = core.data_models.PodiumPackageAsset \
                .where('podium_package_chapter_id', db_chapter.id) \
                .first()

            podbook_episode_chapter_asset = core.data_models.PodiumPackageAsset()
            podbook_episode_chapter_asset.podium_package_id = self.id
            podbook_episode_chapter_asset.type = "podbook_episode_chapter"
            podbook_episode_chapter_asset.format = "text"
            podbook_episode_chapter_asset.accepted_variant = True
            podbook_episode_chapter_asset.title = title
            podbook_episode_chapter_asset.content = content
            podbook_episode_chapter_asset.start_seconds = db_chapter.start
            podbook_episode_chapter_asset.end_seconds = db_chapter.end
            podbook_episode_chapter_asset.derived_from_id = derived_from_asset.id
            podbook_episode_chapter_asset.save()
            podbook_episode_chapter_asset = podbook_episode_chapter_asset.fresh()
        
            podbook_episode_chapter_notes_asset = core.data_models.PodiumPackageAsset()
            podbook_episode_chapter_notes_asset.podium_package_id = self.id
            podbook_episode_chapter_notes_asset.type = "podbook_episode_chapter_notes"
            podbook_episode_chapter_notes_asset.format = "text"
            podbook_episode_chapter_notes_asset.accepted_variant = True
            podbook_episode_chapter_notes_asset.title = title
            podbook_episode_chapter_notes_asset.content = content_notes
            podbook_episode_chapter_notes_asset.start_seconds = db_chapter.start
            podbook_episode_chapter_notes_asset.end_seconds = db_chapter.end
            podbook_episode_chapter_notes_asset.derived_from_id = podbook_episode_chapter_asset.id
            podbook_episode_chapter_notes_asset.save()

    def generate_podbook_episode_chapter_notes(self, content):
        prompt = f"""
[START Content]
{content}
[END Content]

Please write a one-paragraph summary of the key information above. The summary paragraph should be written in the form of a narrative. The paragraph should read like an abridged version of the above content. Remember to include names. 

After the summary, provide a bulleted list of key take-aways from the content above.
Start the bulleted list of key takeaways with "Key Take-aways:"
Use '-' for the bullets.
"""

        content_notes = core.inference.gpt_chat_api_single_prompt(prompt, 'gpt-3.5-turbo-16k', max_tokens=1536, temperature=1.0, top_p=1.0)
        return content_notes

    # ----------------------------------------------------------------------------------------------------------------------------
    # Utility
    # ----------------------------------------------------------------------------------------------------------------------------
    def remove_podbook_thoughts(self, text):
        lines = text.split("\n")
        result = []
        for line in lines:
            if not line.startswith("[") \
            and not line.endswith("]"):
                result.append(line)
        return "\n".join(result)

    def get_podbook_title_and_content(self, text):
        title = ""
        content = ""
        lines = text.split("\n")
        for line in lines:
            if line.startswith("SECTION"):
                title = line.replace("SECTION", "").strip()
                title = title.replace("TITLE", "").strip()
                if title[0] == ":":
                    title = title[1:]
            else:
                content += line + "\n"

        return title.strip(), self.replace_podbook_consecutive_newlines(content).strip()
    
    def replace_podbook_consecutive_newlines(self, text):
      return re.sub(r'\n+', '\n', text)
    
    def get_podbook_transcript_text_with_speakers_from_sentences(self, sentences):
        transcript_text = ""
        current_speaker = None
        for sentence in sentences:
            print(sentence)
            for element in sentence['word_elements']:
                if element['speaker_name'] != current_speaker:
                    current_speaker = element['speaker_name']
                    transcript_text += f"\n\n{element['speaker_name']}:\n"
                if element['value'] in [".", "?", "!", ","] and len(transcript_text) > 0:
                    # remove the space before the punctuation
                    transcript_text = transcript_text[:-1]
                transcript_text += f"{element['value']} "

        return transcript_text
    
    def generate_podbook_files(self):
        podcast_title = ''
        if self.podcast_id:
            podcast = core.data_models.Podcast.find(self.podcast_id)
            if podcast and podcast.title:
                podcast_title = podcast.title

        episode_title = self.get_accepted_asset('titles').content
        if self.podcast_episode_id:
            podcast_episode = core.data_models.PodcastEpisode.find(self.podcast_episode_id)
            if podcast_episode and podcast_episode.title:
                episode_title = podcast_episode.title

        doc = self.get_default_podbook_document()
        notes_doc = self.get_default_podbook_document()
        txt_file_string = ""
        notes_txt_file_string = ""
        file_completed = True

        p = doc.add_paragraph()
        p.add_run(f"\n{podcast_title}", style='book_title_medium').bold = True
        p = doc.add_paragraph()
        p.add_run(f"\n{episode_title}", style='book_title_medium')
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        p = notes_doc.add_paragraph()
        p.add_run(f"\n{podcast_title}", style='book_title_medium').bold = True
        p = notes_doc.add_paragraph()
        p.add_run(f"\n{episode_title}", style='book_title_medium')
        p = notes_doc.add_paragraph()
        p.add_run(f"\n(Quick Notes)", style='book_title_medium')
        p = notes_doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        txt_file_string += f"{podcast_title}\n\n"
        txt_file_string += f"{episode_title}\n\n"
        notes_txt_file_string += f"{podcast_title}\n\n"
        notes_txt_file_string += f"{episode_title}\n\n"

        podbook_episode_chapter_assets = core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'podbook_episode_chapter') \
            .order_by('start_seconds', 'asc') \
            .get()
        
        for index, podbook_episode_chapter_asset in enumerate(podbook_episode_chapter_assets):
            try:
                podbook_episode_chapter_notes_asset = core.data_models.PodiumPackageAsset \
                    .where('podium_package_id', self.id) \
                    .where('type', 'podbook_episode_chapter_notes') \
                    .where('derived_from_id', podbook_episode_chapter_asset.id) \
                    .first()
                
                start_time = int(podbook_episode_chapter_asset.start_seconds)
                if start_time <= 5:
                    start_time = 0
                
                p = doc.add_paragraph()
                p.add_run(f"Section {num2words(index + 1)} ", style='toc_chapter_large').bold = True
                p.add_run(f" {self.format_podbook_seconds(start_time)}\n", style='toc_chapter_large')
                p.add_run(f"{podbook_episode_chapter_asset.title}\n", style='toc_chapter_large')

                p = notes_doc.add_paragraph()
                p.add_run(f"Section {num2words(index + 1)} ", style='toc_chapter_large').bold = True
                p.add_run(f" {self.format_podbook_seconds(start_time)}\n", style='toc_chapter_large')
                p.add_run(f"{podbook_episode_chapter_notes_asset.title}\n", style='toc_chapter_large')
            except Exception as e:
                file_completed = False
                print(f"Error generating book adaptation for package {self.id} asset {podbook_episode_chapter_asset.id}")
                print(e)

        new_section = notes_doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.header.is_linked_to_previous = False
        run = new_section.header.paragraphs[0].add_run(f"", style='page_number')
        new_section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.add_podbook_page_number(run)
                
        for index, podbook_episode_chapter_asset in enumerate(podbook_episode_chapter_assets):
            try:
                podbook_episode_chapter_notes_asset = core.data_models.PodiumPackageAsset \
                    .where('podium_package_id', self.id) \
                    .where('type', 'podbook_episode_chapter_notes') \
                    .where('derived_from_id', podbook_episode_chapter_asset.id) \
                    .first()

                start_time = int(podbook_episode_chapter_asset.start_seconds)
                if start_time <= 5:
                    start_time = 0

                new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.header.is_linked_to_previous = False
                run = new_section.header.paragraphs[0].add_run(f"", style='page_number')
                new_section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self.add_podbook_page_number(run)

                p = doc.add_paragraph()
                p.add_run(f"\n({self.format_podbook_seconds(start_time)})", style='paragraph_content')
                p.add_run(f"\n{podbook_episode_chapter_asset.title}\n", style='subsection_heading').bold = True
                
                p = notes_doc.add_paragraph()
                p.add_run(f"\n({self.format_podbook_seconds(start_time)})", style='paragraph_content')
                p.add_run(f"\n{podbook_episode_chapter_notes_asset.title}\n", style='subsection_heading').bold = True
                
                txt_file_string += f"\n\n({self.format_podbook_seconds(start_time)})\n{podbook_episode_chapter_asset.title}\n\n"
                notes_txt_file_string += f"\n\n({self.format_podbook_seconds(start_time)})\n{podbook_episode_chapter_notes_asset.title}\n\n"

                formatted_content_notes = podbook_episode_chapter_notes_asset.content.replace('- ', '\n- ')
                p = notes_doc.add_paragraph()
                p.add_run(f"{formatted_content_notes}", style='paragraph_content')
                notes_txt_file_string += f"{formatted_content_notes}\n\n"

                for paragraph in podbook_episode_chapter_asset.content.split('\n'):
                    p = doc.add_paragraph()
                    p.add_run(f"{paragraph}", style='paragraph_content')
                    txt_file_string += f"{paragraph}\n\n"
            except Exception as e:
                file_completed = False
                print(f"Error generating book adaptation for package {self.id} asset {podbook_episode_chapter_asset.id}")
                print(e)

        if file_completed:
            file_name = f"{self.id}_{'_'.join(podcast_title.split(' ')[:10])}"
            notes_file_name = f"{self.id}_notes_{'_'.join(podcast_title.split(' ')[:10])}"
            # replace any invalid chars
            file_name = file_name.replace('/', '-')
            notes_file_name = notes_file_name.replace('/', '-')

            if not os.path.exists('podbooks'):
                os.makedirs('podbooks')
            
            doc.save(f"podbooks/{file_name}.docx")
            notes_doc.save(f"podbooks/{notes_file_name}.docx")

            with open(f"podbooks/{file_name}.txt", 'w') as f:
                f.write(txt_file_string)

            with open(f"podbooks/{notes_file_name}.txt", 'w') as f:
                f.write(notes_txt_file_string)

    def get_default_podbook_document(self):
        doc = Document()

        obj_styles = doc.styles
        obj_charstyle = obj_styles.add_style('book_title_small', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(22)
        obj_font.name = 'Calibri'
    
        obj_styles = doc.styles
        obj_charstyle = obj_styles.add_style('book_title_medium', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(32)
        obj_font.name = 'Calibri'

        obj_styles = doc.styles
        obj_charstyle = obj_styles.add_style('book_title_large', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(62)
        obj_font.name = 'Calibri'
        
        obj_charstyle = obj_styles.add_style('toc_chapter_small', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(12)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('toc_chapter_large', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(16)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('chapter_heading_small', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(18)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('chapter_heading_large', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(28)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('chapter_heading_xlarge', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(39)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('subsection_heading', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(17)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('paragraph_content', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(15)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('small_paragraph_content', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(13)
        obj_font.name = 'Calibri'

        obj_charstyle = obj_styles.add_style('page_number', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(12)
        obj_font.name = 'Calibri'

        return doc

    def add_podbook_page_number(self, run):
        fldChar1 = self.create_podbook_element('w:fldChar')
        self.create_podbook_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_podbook_element('w:instrText')
        self.create_podbook_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_podbook_element('w:fldChar')
        self.create_podbook_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def format_podbook_seconds(self, seconds):
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return "{:02d}:{:02d}:{:02d}".format(h, m, s)
    
    def create_podbook_element(self, name):
        return OxmlElement(name)
    
    def create_podbook_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def generate_expo_file(self):
        podcast_title = ''
        if self.podcast_id:
            podcast = core.data_models.Podcast.find(self.podcast_id)
            if podcast and podcast.title:
                podcast_title = podcast.title

        doc = self.get_default_podbook_document()

        p = doc.add_paragraph()
        p.add_run(f"\nCreate more with Podium.", style='chapter_heading_xlarge').bold = True
        p = doc.add_paragraph()
        p.add_run(f"\nLet Podium do the heavy lifting for your next episode!", style='book_title_medium')
        p = doc.add_paragraph()
        p.add_run(f"\nHere's an example of the Titles, Shownotes, Chapters, Highlights and Keywords that Podium created for this episode - plus, you can create anything with PodiumGPT!", style='book_title_small')
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        titles = []
        accepted_title = self.get_accepted_asset('titles')
        titles.append(accepted_title.content)
        variants = core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'titles') \
            .where('parent_id',accepted_title.id) \
            .get()
        
        for variant in variants:
            titles.append(variant.content)

        p = doc.add_paragraph()
        p.add_run(f"\nTitle Options\n", style='subsection_heading').bold = True
        for title in titles:
            p = doc.add_paragraph()
            p.add_run(f"{title}\n", style='paragraph_content')

        accepted_keywords = self.get_accepted_asset('keywords')
        p = doc.add_paragraph()
        p.add_run(f"\nKeywords\n", style='subsection_heading').bold = True
        p = doc.add_paragraph()
        p.add_run(f"{accepted_keywords.content}\n", style='paragraph_content')


        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        shownotes_summaries = []
        accepted_shownotes_summary = self.get_accepted_asset('show_notes_summary')
        shownotes_summaries.append(accepted_shownotes_summary.content)
        variants = core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'show_notes_summary') \
            .where('parent_id',accepted_shownotes_summary.id) \
            .get()
        for variant in variants:
            shownotes_summaries.append(variant.content)

        p = doc.add_paragraph()
        p.add_run(f"Shownotes Options\n", style='subsection_heading').bold = True
        for shownotes_summary in shownotes_summaries:
            p = doc.add_paragraph()
            p.add_run(f"{shownotes_summary}\n", style='small_paragraph_content')
            p.add_run(f"\n___\n", style='small_paragraph_content')


        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        p = doc.add_paragraph()
        p.add_run(f"Chapters\n", style='subsection_heading').bold = True

        chapters = []
        accepted_chapters = core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'chapter') \
            .where('accepted_variant', True) \
            .order_by('start_seconds', 'asc') \
            .get()
    
        for chapter in accepted_chapters:
            chapter_variant = core.data_models.PodiumPackageAsset \
                .where('podium_package_id', self.id) \
                .where('type', 'chapter') \
                .where('parent_id', chapter.id) \
                .first()
            p = doc.add_paragraph()
            p.add_run(f"({self.format_podbook_seconds(int(chapter.start_seconds))}) {chapter.title}\n\n", style='paragraph_content').bold = True
            p.add_run(f"(Short Description)\n{chapter_variant.content}\n\n", style='small_paragraph_content')
            p.add_run(f"(Long Description)\n{chapter.content}\n\n", style='small_paragraph_content')
        

        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        p = doc.add_paragraph()
        p.add_run(f"Highlights\n", style='subsection_heading').bold = True

        accepted_highlights = core.data_models.PodiumPackageAsset \
            .where('podium_package_id', self.id) \
            .where('type', 'highlight') \
            .where('accepted_variant', True) \
            .order_by('start_seconds', 'asc') \
            .get()
        
        transcript_file = self.transcript_file()
        transcript_words = transcript_file.get_words_and_punct()
        
        for highlight in accepted_highlights:
            p = doc.add_paragraph()
            p.add_run(f"\n({self.format_podbook_seconds(int(highlight.start_seconds))} - {self.format_podbook_seconds(int(highlight.end_seconds))}) {highlight.title}", style='paragraph_content').bold = True
            
            highlight_transcript_words = core.data_models.PodiumPackageTranscriptFile.get_words_between_start_end(transcript_words, float(highlight.start_seconds), float(highlight.end_seconds))
            elements = []
            previous_speaker_id = None
            for index, item in enumerate(highlight_transcript_words): 
                if item['speaker_id'] != previous_speaker_id and 'start' in item:
                    if len(elements) > 0:
                        monologue = "".join(elements)
                        p = doc.add_paragraph()
                        p.add_run(f"{monologue}", style='small_paragraph_content')

                    elements = []
                    elements.append(item['value'])
                    previous_speaker_id = item['speaker_id']
                    
                    p = doc.add_paragraph()
                    p.add_run(f"({self.format_podbook_seconds(int(item['start']))}) {item['speaker_name']}", style='small_paragraph_content')
                else:
                    elements.append(item['value'])

            monologue = "".join(elements)
            p = doc.add_paragraph()
            p.add_run(f"{monologue}", style='small_paragraph_content')
        

        file_name = f"{self.id}_expo_{'_'.join(podcast_title.split(' ')[:10])}"
        doc.save(f"podbooks/{file_name}.docx")













